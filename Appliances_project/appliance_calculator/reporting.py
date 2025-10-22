import datetime
from typing import List, Dict, Union
from docx import Document
from openpyxl import Workbook 

# Сохранение в .doc и .xls
class Report:
    """
    Абстрактный базовый класс для генерации отчетов.
    """
    def __init__(self, filename: str):
        if not isinstance(filename, str) or not filename:
            raise ValueError("Имя файла должно быть непустой строкой.")
        self._filename = filename

    @property
    def filename(self) -> str:
        return self._filename

    # 2 dunder-метода
    def __str__(self) -> str:
        return f"Отчет (файл: {self.filename})"

    def __repr__(self) -> str:
        return f"Report(filename='{self.filename}')"

    def save_report(self, data: Dict[str, Union[float, List[Dict[str, Union[str, float]]]]]):
        """
        Метод для сохранения данных в файл. Должен быть реализован в подклассах.
        """
        raise NotImplementedError("Подклассы должны реализовать этот метод.")

class DocReport(Report):
    """
    Класс для сохранения отчета в формате .doc.
    """
    def save_report(self, data: Dict[str, Union[float, List[Dict[str, Union[str, float]]]]]):
        document = Document()
        document.add_heading('Отчет о потреблении электроэнергии', 0)

        document.add_paragraph(f"Дата генерации: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_paragraph(f"Общее потребление: {data['total_consumption_kwh']} кВт*ч")
        document.add_paragraph(f"Общая стоимость: {data['total_cost']} $.") # $. - условные единицы

        document.add_heading('Детализация по приборам', level=1)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Прибор'
        hdr_cells[1].text = 'Потребление (кВт*ч)'
        hdr_cells[2].text = 'Стоимость ($.)'

        for item in data['details']:
            row_cells = table.add_row().cells
            row_cells[0].text = item['name']
            row_cells[1].text = str(item['consumption_kwh'])
            row_cells[2].text = str(item['cost'])

        for cell in table.columns[0].cells: cell.width = 1500000
        for cell in table.columns[1].cells: cell.width = 500000
        for cell in table.columns[2].cells: cell.width = 500000

        try:
            document.save(f"{self.filename}.docx")
            print(f"Отчет успешно сохранен в файл: {self.filename}.docx")
        except Exception as e:
            print(f"Ошибка при сохранении DOCX файла: {e}")


class XlsReport(Report):
    """
    Класс для сохранения отчета в формате .xls.
    """
    def save_report(self, data: Dict[str, Union[float, List[Dict[str, Union[str, float]]]]]):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Энергопотребление"

        sheet.append(["Отчет о потреблении электроэнергии", datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
        sheet.append([]) # Пустая строка

        sheet.append(["Общее потребление (кВт*ч):", data['total_consumption_kwh']])
        sheet.append(["Общая стоимость ($.):", data['total_cost']])
        sheet.append([]) # Пустая строка

        sheet.append(["Прибор", "Потребление (кВт*ч)", "Стоимость ($.)"])
        for item in data['details']:
            sheet.append([item['name'], item['consumption_kwh'], item['cost']])

        try:
            workbook.save(f"{self.filename}.xlsx")
            print(f"Отчет успешно сохранен в файл: {self.filename}.xlsx")
        except Exception as e:
            print(f"Ошибка при сохранении XLSX файла: {e}")
